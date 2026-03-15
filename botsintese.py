#!/usr/bin/env python3
"""
BotSíntese v3.0 - Síntese Processual Automatizada
Extração, organização e síntese factual de processos judiciais.

Suporta:
- Processamento local (Ollama/Llama)
- Cloud: Google Gemini, Anthropic Claude, OpenAI GPT, xAI Grok

Autor: Gerado por Claude (Anthropic)
Versão: 2.0.0
"""

import os
import sys
import re
import json
import time
import hashlib
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, field
from typing import Optional, List, Dict, Tuple
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading
import builtins

# Dependências externas
try:
    import requests
    import yaml
    from PyPDF2 import PdfReader
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError as e:
    print(f"Erro: Dependência não encontrada - {e}")
    print("Execute: pip install requests pyyaml PyPDF2 python-docx --break-system-packages")
    sys.exit(1)

# ============================================================================
# CONFIGURAÇÕES
# ============================================================================

def print_seguro(*args, **kwargs):
    """Faz print sem quebrar quando o terminal nao suporta Unicode."""
    try:
        builtins.print(*args, **kwargs)
    except UnicodeEncodeError:
        sep = kwargs.get("sep", " ")
        end = kwargs.get("end", "\n")
        arquivo = kwargs.get("file", sys.stdout)
        flush = kwargs.get("flush", False)
        encoding = getattr(arquivo, "encoding", None) or "utf-8"
        texto = sep.join(str(arg) for arg in args) + end
        arquivo.write(texto.encode(encoding, errors="replace").decode(encoding, errors="replace"))
        if flush:
            arquivo.flush()


print = print_seguro


@dataclass
class Config:
    """Configurações do BotSíntese"""
    # APIs Cloud
    api_anthropic: str = ""
    api_openai: str = ""
    api_google: str = ""
    api_xai: str = ""
    
    # Ollama (local)
    ollama_host: str = "http://localhost:11434"
    modelo_local: str = "llama3.1:8b-instruct-q4_K_M"
    
    # Modo padrão
    modo_padrao: str = "google"  # local, google, anthropic, openai, xai
    
    # Processamento - chunks maiores para aproveitar contexto do Gemini
    chunk_size_local: int = 6000   # ~24k chars para Llama 8B
    chunk_size_cloud: int = 200000  # ~800k chars para Gemini (tem 1M contexto)
    chars_per_token: float = 4.0


def carregar_config(pasta_script: Path) -> Config:
    """Carrega configurações do arquivo YAML"""
    config = Config()
    
    # Procura config na pasta do script
    config_file = pasta_script / "botsintese_config.yaml"
    
    if config_file.exists():
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                dados = yaml.safe_load(f) or {}
                
            # APIs
            apis = dados.get('apis', {})
            config.api_anthropic = apis.get('anthropic', '')
            config.api_openai = apis.get('openai', '')
            config.api_google = apis.get('google', '')
            config.api_xai = apis.get('xai', '')
            
            # Ollama
            ollama = dados.get('ollama', {})
            config.ollama_host = ollama.get('host', config.ollama_host)
            config.modelo_local = ollama.get('modelo', config.modelo_local)
            
            # Geral
            config.modo_padrao = dados.get('modo_padrao', 'local')
            
        except Exception as e:
            print(f"Aviso: Erro ao carregar config: {e}")
    
    return config


def salvar_config(config: Config, pasta_script: Path):
    """Salva configurações no arquivo YAML"""
    config_file = pasta_script / "botsintese_config.yaml"
    
    dados = {
        'apis': {
            'anthropic': config.api_anthropic,
            'openai': config.api_openai,
            'google': config.api_google,
            'xai': config.api_xai,
        },
        'ollama': {
            'host': config.ollama_host,
            'modelo': config.modelo_local,
        },
        'modo_padrao': config.modo_padrao,
    }
    
    with open(config_file, 'w', encoding='utf-8') as f:
        yaml.dump(dados, f, allow_unicode=True, default_flow_style=False)


# ============================================================================
# DETECÇÃO DE SISTEMA PROCESSUAL
# ============================================================================

@dataclass
class EventoProcessual:
    """Representa um evento/documento do processo"""
    data: str
    tipo: str
    descricao: str
    pagina_inicio: Optional[int] = None
    pagina_fim: Optional[int] = None
    conteudo: str = ""


@dataclass
class DadosProcesso:
    """Dados extraídos do processo"""
    numero: str = ""
    classe: str = ""
    vara: str = ""
    comarca: str = ""
    valor_causa: str = ""
    data_distribuicao: str = ""
    assunto: str = ""
    
    partes: List[Dict] = field(default_factory=list)
    eventos: List[EventoProcessual] = field(default_factory=list)
    
    sistema: str = "generico"  # pje, eproc, saj, projudi, generico


def detectar_sistema(texto: str) -> str:
    """Detecta qual sistema processual gerou o PDF"""
    texto_lower = texto[:10000].lower()
    
    if "pje - processo judicial eletrônico" in texto_lower or "pje.tjmg" in texto_lower:
        return "pje"
    elif "página de separação" in texto_lower and "evento" in texto_lower:
        return "eproc"
    elif "projudi" in texto_lower:
        return "projudi"
    elif any(t in texto_lower for t in ["saj", "esaj", "e-saj", "saj/pg5", "sajpg", "portal de serviços e-saj"]):
        return "saj"
    elif any(t in texto_lower for t in ["foro de", "foro central", "foro regional", "tribunal de justiça do estado de são paulo", "tjsp"]):
        return "saj"  # TJSP usa SAJ
    else:
        return "generico"


def extrair_dados_pje(texto: str) -> DadosProcesso:
    """Extrai dados de PDF do sistema PJe"""
    dados = DadosProcesso(sistema="pje")
    
    # Número do processo
    match = re.search(r'Número:\s*([\d.-]+)', texto)
    if match:
        dados.numero = match.group(1).strip()
    
    # Classe
    match = re.search(r'Classe:\s*\[?\w*\]?\s*([^\n]+)', texto)
    if match:
        dados.classe = match.group(1).strip()
    
    # Órgão julgador
    match = re.search(r'Órgão julgador:\s*([^\n]+)', texto)
    if match:
        dados.vara = match.group(1).strip()
    
    # Valor da causa
    match = re.search(r'Valor da causa:\s*R?\$?\s*([\d.,]+)', texto)
    if match:
        dados.valor_causa = f"R$ {match.group(1).strip()}"
    
    # Data distribuição
    match = re.search(r'(?:Última )?[Dd]istribuição\s*:?\s*(\d{2}/\d{2}/\d{4})', texto)
    if match:
        dados.data_distribuicao = match.group(1)
    
    # Assunto
    match = re.search(r'Assuntos?:\s*([^\n]+)', texto)
    if match:
        dados.assunto = match.group(1).strip()
    
    # Partes - busca na tabela
    # Padrão: NOME (TIPO) seguido opcionalmente de ADVOGADO
    partes_pattern = r'([A-ZÁÉÍÓÚÇÃÕ][A-ZÁÉÍÓÚÇÃÕ\s]+)\s*\((AUTOR|RÉU|RÉ|REQUERENTE|REQUERIDO|APELANTE|APELADO)[^)]*\)'
    for match in re.finditer(partes_pattern, texto[:3000]):
        nome = match.group(1).strip()
        polo = match.group(2).strip()
        if len(nome) > 3:
            dados.partes.append({
                'nome': nome,
                'polo': 'Autor' if polo in ['AUTOR', 'REQUERENTE', 'APELANTE'] else 'Réu',
                'advogado': ''
            })
    
    # Eventos/Documentos - busca na tabela do PJe
    # Padrão: ID | Data | Documento | Tipo
    evento_pattern = r'(\d{2}/\d{2}/\d{4}\s+\d{2}:\d{2})\s+([^\n]+?)\s+(Petição|Contestação|Sentença|Despacho|Decisão|Certidão|Intimação|Citação|Manifestação|Acórdão|Recurso|Laudo|Impugnação|Réplica)[^\n]*'
    for match in re.finditer(evento_pattern, texto, re.IGNORECASE):
        data = match.group(1).split()[0]  # Só a data, sem hora
        descricao = match.group(2).strip()
        tipo = match.group(3).strip()
        
        dados.eventos.append(EventoProcessual(
            data=data,
            tipo=tipo,
            descricao=descricao[:100]
        ))
    
    return dados


def extrair_dados_eproc(texto: str) -> DadosProcesso:
    """Extrai dados de PDF do sistema e-Proc"""
    dados = DadosProcesso(sistema="eproc")
    
    # Número do processo
    match = re.search(r'Processo:\s*([\d.-]+)', texto)
    if match:
        dados.numero = match.group(1).strip()
    
    # Eventos - padrão e-Proc com página de separação
    evento_pattern = r'Evento\s+(\d+).*?Data:\s*(\d{2}/\d{2}/\d{4})[^\n]*.*?(?:Tipo|Documento):\s*([^\n]+)'
    for match in re.finditer(evento_pattern, texto, re.DOTALL | re.IGNORECASE):
        numero = match.group(1)
        data = match.group(2)
        tipo = match.group(3).strip()
        
        dados.eventos.append(EventoProcessual(
            data=data,
            tipo=tipo,
            descricao=f"Evento {numero}"
        ))
    
    return dados
    return dados


def extrair_dados_saj(texto: str) -> DadosProcesso:
    """Extrai dados de PDF do sistema SAJ (TJSP/TJAK/etc)"""
    dados = DadosProcesso(sistema="saj")
    texto_inicio = texto[:5000]  # Cabeçalho costuma estar no início
    
    # Número do processo
    match = re.search(r'Processo\s*n[º°]:?\s*([\d.-]+)', texto_inicio, re.IGNORECASE)
    if match:
        dados.numero = match.group(1).strip()
    
    # Classe - Assunto (Padrão SAJ: "Classe - Assunto: Execução... - Nota Promissória")
    # Tenta pegar a linha completa primeiro
    match = re.search(r'Classe\s*-\s*Assunto:?\s*([^\n]+)', texto_inicio, re.IGNORECASE)
    if match:
        conteudo = match.group(1).strip()
        if " - " in conteudo:
            partes_classe = conteudo.split(" - ", 1)
            dados.classe = partes_classe[0].strip()
            dados.assunto = partes_classe[1].strip()
        else:
            dados.classe = conteudo
    else:
        # Fallback para "Classe:" isolado
        match = re.search(r'Classe:?\s*([^\n]+)', texto_inicio, re.IGNORECASE)
        if match:
            dados.classe = match.group(1).strip()
            
    # Foro / Comarca
    match = re.search(r'Foro\s*(?:de|da|do)?\s*([^\n]+)', texto_inicio, re.IGNORECASE)
    if not match:
        match = re.search(r'Comarca\s*(?:de|da|do)?\s*([^\n]+)', texto_inicio, re.IGNORECASE)
    if match:
        dados.comarca = match.group(1).strip()
        
    # Vara
    match = re.search(r'(\d+ª\s*Vara\s*[^\n]+)', texto_inicio, re.IGNORECASE)
    if match:
        dados.vara = match.group(1).strip()
        
    # Data de distribuição (muitas vezes aparece como "Distribuição:")
    match = re.search(r'Distribuição:?\s*(\d{2}/\d{2}/\d{4})', texto_inicio, re.IGNORECASE)
    if match:
        dados.data_distribuicao = match.group(1)
        
    # Juiz
    match = re.search(r'Juiz\(a\)\s*de\s*Direito:?\s*Dr\(a\)\.\s*([^\n]+)', texto_inicio, re.IGNORECASE)
    # Não temos campo Juiz no DadosProcesso, mas ajuda a confirmar que é cabeçalho
    
    # Partes (Exequente / Executado / Requerente / Requerido)
    # Padrão SAJ: "Exequente: Nome..."
    #             "Executado: Nome..."
    polos_map = {
        'Exequente': 'Autor', 'Requerente': 'Autor', 'Autor': 'Autor', 'Embargante': 'Autor',
        'Executado': 'Réu', 'Requerido': 'Réu', 'Réu': 'Réu', 'Embargado': 'Réu'
    }
    
    for label, polo_norm in polos_map.items():
        pattern = rf'{label}:?\s*([^\n]+)'
        for match in re.finditer(pattern, texto_inicio, re.IGNORECASE):
            nome = match.group(1).strip()
            # Evita pegar texto processual que venha depois (ex: "Exequente: Nome do cara. Vistos...")
            if "." in nome:
                nome = nome.split(".")[0]
            
            if len(nome) > 3 and "juiz" not in nome.lower():
                dados.partes.append({
                    'nome': nome.strip(),
                    'polo': polo_norm
                })

    return dados

def extrair_dados_generico(texto: str) -> DadosProcesso:
    """Extrai dados de PDF sem sistema identificado"""
    dados = DadosProcesso(sistema="generico")
    
    # Usa primeiros 10000 chars para dados estruturados
    texto_inicio = texto[:10000]
    
    # Tenta encontrar número de processo em vários formatos
    patterns_processo = [
        r'(\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4})',  # CNJ
        r'Processo\s*(?:n[ºo.]?)?\s*([\d./-]+)',
        r'Autos\s*(?:n[ºo.]?)?\s*([\d./-]+)',
    ]
    
    for pattern in patterns_processo:
        match = re.search(pattern, texto_inicio)
        if match:
            dados.numero = match.group(1).strip()
            break
    
    # Classe processual
    patterns_classe = [
        r'[Cc]lasse[:\s]+([A-Z][^\n]{3,60})',  # Exige Maiúscula inicial e max 60 chars
        r'[Aa]ção\s+de\s+([A-Z][^\n]{3,60})',
        r'[Tt]ipo\s+de\s+[Aa]ção[:\s]+([^\n]{3,60})',
    ]
    for pattern in patterns_classe:
        match = re.search(pattern, texto_inicio)
        if match:
            # Limpeza extra
            valor = match.group(1).strip().rstrip('.')
            if len(valor) > 3 and "..." not in valor and not valor.isdigit():
                 dados.classe = valor
                 break
    
    # Vara e Foro
    patterns_vara = [
        r'(\d+[ªa]?\s*[Vv]ara\s+[^\n]{3,60})',
        r'[Vv]ara[:\s]+([^\n]{5,60})',
        r'([Jj]uizado\s+[Ee]special\s+[^\n]{3,60})',
    ]
    for pattern in patterns_vara:
        match = re.search(pattern, texto_inicio)
        if match:
            dados.vara = match.group(1).strip().rstrip('.')
            break
    
    # Comarca
    patterns_comarca = [
        r'[Cc]omarca\s+(?:de\s+)?([A-Z][^\n]{3,40})',  # Maiúscula e curto
        r'[Ff]oro\s+(?:(?:da|de|do)\s+)?(?:[Cc]omarca\s+(?:de\s+)?)?([A-Z][^\n]{3,40})',
    ]
    for pattern in patterns_comarca:
        match = re.search(pattern, texto_inicio)
        if match:
            dados.comarca = match.group(1).strip().rstrip('.')
            break
    
    # Data de distribuição
    patterns_dist = [
        r'[Dd]istribui[çc][ãa]o[:\s]*(\d{2}/\d{2}/\d{4})',
        r'[Dd]istribuíd[oa]\s+em[:\s]*(\d{2}/\d{2}/\d{4})',
        r'[Dd]ata\s+de\s+[Dd]istribui[çc][ãa]o[:\s]*(\d{2}/\d{2}/\d{4})',
    ]
    for pattern in patterns_dist:
        match = re.search(pattern, texto_inicio)
        if match:
            dados.data_distribuicao = match.group(1)
            break
    
    # Assunto
    patterns_assunto = [
        r'[Aa]ssunto[:\s]+([A-Z][^\n]{3,80})',
        r'[Aa]ssunto\s+[Pp]rincipal[:\s]+([A-Z][^\n]{3,80})',
    ]
    for pattern in patterns_assunto:
        match = re.search(pattern, texto_inicio)
        if match:
            dados.assunto = match.group(1).strip().rstrip('.')
            break
    
    # Valor da causa
    match = re.search(r'[Vv]alor\s+(?:da\s+)?[Cc]ausa[:\s]*R?\$?\s*([\d.,]+)', texto)
    if match:
        dados.valor_causa = f"R$ {match.group(1)}"
    
    return dados


# ============================================================================
# EXTRAÇÃO DE PDF
# ============================================================================

def extrair_texto_pdf(caminho: Path) -> Tuple[str, int]:
    """Extrai texto de um PDF. Retorna (texto, num_paginas)"""
    texto = []
    num_paginas = 0
    
    try:
        reader = PdfReader(str(caminho))
        num_paginas = len(reader.pages)
        
        for i, pagina in enumerate(reader.pages):
            texto_pagina = pagina.extract_text() or ""
            if texto_pagina.strip():
                texto.append(f"\n[PÁGINA {i+1}]\n{texto_pagina}")
    except Exception as e:
        print(f"Erro ao ler {caminho.name}: {e}")
        return "", 0
    
    return "\n".join(texto), num_paginas


def dividir_em_chunks(texto: str, config: Config, modo: str = "local") -> List[str]:
    """Divide texto em chunks para processamento"""
    # Usa chunk maior para cloud (tem mais contexto)
    if modo in ["google", "anthropic", "openai", "xai"]:
        chunk_size = config.chunk_size_cloud
    else:
        chunk_size = config.chunk_size_local
    
    max_chars = int(chunk_size * config.chars_per_token)
    chunks = []
    
    # Divide por páginas
    paginas = re.split(r'\n\[PÁGINA \d+\]\n', texto)
    paginas = [p for p in paginas if p.strip()]
    
    chunk_atual = ""
    for pagina in paginas:
        if len(chunk_atual) + len(pagina) < max_chars:
            chunk_atual += "\n" + pagina
        else:
            if chunk_atual.strip():
                chunks.append(chunk_atual.strip())
            # Se uma única página for maior que max_chars, divide ela
            if len(pagina) > max_chars:
                # Divide em partes menores
                for i in range(0, len(pagina), max_chars):
                    chunks.append(pagina[i:i+max_chars])
                chunk_atual = ""
            else:
                chunk_atual = pagina
    
    if chunk_atual.strip():
        chunks.append(chunk_atual.strip())
    
    return chunks if chunks else [texto[:max_chars]]


# ============================================================================
# PROMPTS
# ============================================================================

PROMPT_EXTRACAO = """Você é um assistente de extração de dados processuais. Analise o texto e extraia APENAS informações factuais, sem fazer análises ou sugestões.

TEXTO DO PROCESSO:
{texto}

Extraia as seguintes informações em formato JSON (retorne APENAS o JSON, sem texto adicional):

{{
    "partes": [
        {{"nome": "Nome completo da parte", "polo": "Autor/Réu"}}
    ],
    "classe_processual": "Tipo da ação (ex: Execução de Título Extrajudicial, Ação de Cobrança, etc.)",
    "vara": "Vara e Foro onde tramita o processo",
    "comarca": "Comarca do processo",
    "data_distribuicao": "dd/mm/aaaa",
    "assunto": "Assunto principal do processo",
    "objeto_acao": "Descrição breve do que se trata a ação (1-2 frases)",
    "resumo_fatos": "Narrativa dos fatos em parágrafos bem formatados, com quebras de linha (\\n\\n) entre parágrafos. Conte a história do processo de forma clara e cronológica.",
    "valores_relevantes": [
        {{"descricao": "...", "valor": "R$ ..."}}
    ],
    "pedidos": ["Pedido 1", "Pedido 2"],
    "decisoes": [
        {{"data": "dd/mm/aaaa", "tipo": "Despacho/Decisão/Sentença", "conteudo": "Resumo do que foi decidido"}}
    ],
    "teses_autor": ["Tese 1", "Tese 2"],
    "teses_reu": ["Tese 1", "Tese 2"],
    "documentos_importantes": [
        {{"tipo": "Petição Inicial/Contestação/Sentença/etc", "data": "dd/mm/aaaa", "parte": "Quem apresentou", "resumo": "Resumo do conteúdo principal do documento"}}
    ],
    "historico_detalhado": [
        {{"data": "dd/mm/aaaa", "evento": "Tipo do evento", "descricao": "O que aconteceu de fato, quem fez, qual o conteúdo resumido"}}
    ],
    "status_atual": "Descreva o status atual baseado no ÚLTIMO despacho/decisão encontrado no texto. Inclua a data e o conteúdo resumido do último ato judicial."
}}

REGRAS IMPORTANTES:
- RETORNE APENAS O JSON, sem explicações antes ou depois
- Extraia APENAS o que está explícito no texto
- NÃO invente informações
- NÃO faça análises ou sugestões jurídicas
- PARTES: Extraia da PETIÇÃO INICIAL. O autor é quem propõe a ação. Os réus são contra quem a ação é proposta. Use os nomes completos das pessoas/empresas, não use termos como "Contestante", "Requerido", etc.
- Datas no formato dd/mm/aaaa
- Se não encontrar uma informação, deixe vazio ou null
- No resumo_fatos, use parágrafos separados por \\n\\n para facilitar leitura
- Em valores_relevantes, inclua APENAS valores diretamente relacionados à causa (valor da causa, valores cobrados, danos pedidos). NÃO inclua capital social de empresas, valor de cotas, salários, etc.
- Em documentos_importantes, foque nas peças processuais principais: petição inicial, contestações, réplicas, decisões, sentenças, laudos
- Em historico_detalhado:
  * Inclua TODOS os eventos relevantes encontrados, cobrindo desde o ajuizamento até o ato mais recente
  * Inclua especificamente: despachos, decisões, citações, intimações, penhoras, avaliações, leilões, expedição de mandados
  * Inclua TODOS os despachos e decisões, mesmo os mais simples como "Cite-se" ou "Dê-se vista"
  * O último evento deve ser o ato mais recente encontrado no texto
  * Seja específico: não "Manifestação", mas "Manifestação do autor sobre citação"
  * PRIORIZE completude — é melhor incluir um evento a mais do que perder um despacho importante
- Em status_atual, baseie-se SEMPRE no último despacho/decisão do processo, não em suposições sobre a fase genérica
- Seja conciso e objetivo"""


PROMPT_CONSOLIDACAO = """Você é um assistente de síntese processual. Consolide as extrações parciais abaixo em um único documento coerente.

EXTRAÇÕES PARCIAIS:
{extracoes}

DADOS JÁ CONHECIDOS DO PROCESSO:
- Número: {numero}
- Sistema: {sistema}
- Eventos já identificados: {num_eventos}

Gere uma consolidação em formato JSON (retorne APENAS o JSON, sem texto adicional):

{{
    "objeto_acao": "Síntese clara do objeto da ação em 1-2 frases",
    "resumo_fatos": "Narrativa cronológica dos fatos, bem formatada com parágrafos separados por \\n\\n. Máximo 600 palavras. Conte a história de forma clara.",
    "partes_consolidadas": [
        {{"nome": "Nome completo", "polo": "Autor/Réu"}}
    ],
    "valores_relevantes": [
        {{"descricao": "...", "valor": "R$ ..."}}
    ],
    "teses_autor": [...],
    "teses_reu": [...],
    "documentos_importantes": [
        {{"tipo": "...", "data": "...", "parte": "...", "resumo": "Resumo detalhado do documento (3-5 frases)"}}
    ],
    "decisoes_importantes": [...],
    "historico_resumido": [
        {{"data": "...", "descricao": "Descrição clara do que aconteceu"}}
    ],
    "status_atual": "..."
}}

REGRAS:
- RETORNE APENAS O JSON, sem explicações antes ou depois
- Elimine redundâncias e informações duplicadas
- Mantenha apenas fatos confirmados
- NÃO faça análises jurídicas ou sugestões
- PARTES: Use nomes completos das pessoas/empresas. NÃO use termos genéricos como "Contestante", "Requerido", "Primeiro Réu". Use os nomes reais.
- Use parágrafos no resumo_fatos (separados por \\n\\n)
- Em valores_relevantes, inclua APENAS: valor da causa, valores contratuais em disputa, valores de danos pedidos. EXCLUA: capital social, cotas de empresas, dados societários
- Em documentos_importantes, faça um resumo útil de cada peça principal
- Seja factual e objetivo"""


# ============================================================================
# FUNÇÕES DE NORMALIZAÇÃO E LIMPEZA (v3.0)
# ============================================================================

INVALID_ESCAPE_RE = re.compile(r'(?<!\\)\\(?!["\\/bfnrtu])')
INVALID_UNICODE_ESCAPE_RE = re.compile(r'(?<!\\)\\u(?![0-9a-fA-F]{4})')
CONTROL_CHAR_RE = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f]')


def extrair_json_candidato(resposta: str) -> str:
    """Isola o JSON principal de uma resposta do modelo."""
    if not resposta:
        return ""
    
    texto = resposta.strip().lstrip("\ufeff")
    texto = re.sub(r'^```(?:json)?\s*', '', texto, flags=re.IGNORECASE)
    texto = re.sub(r'\s*```$', '', texto)
    
    inicio = texto.find('{')
    fim = texto.rfind('}')
    if inicio >= 0 and fim > inicio:
        return texto[inicio:fim+1].strip()
    
    return texto


def gerar_candidatos_json(json_str: str) -> List[str]:
    """Gera versoes progressivamente mais tolerantes do JSON bruto."""
    candidatos = []
    
    def adicionar(valor: str):
        valor = (valor or "").strip()
        if valor and valor not in candidatos:
            candidatos.append(valor)
    
    base = (json_str or "").strip().lstrip("\ufeff")
    base = (
        base.replace("\u201c", '"')
            .replace("\u201d", '"')
            .replace("\u2018", "'")
            .replace("\u2019", "'")
    )
    adicionar(base)
    
    sem_controle = CONTROL_CHAR_RE.sub(' ', base)
    adicionar(sem_controle)
    
    corrigido = re.sub(r',\s*([}\]])', r'\1', sem_controle)
    corrigido = re.sub(r'}\s*"', '}, "', corrigido)
    corrigido = re.sub(r']\s*"', '], "', corrigido)
    adicionar(corrigido)
    
    corrigido_escapes = INVALID_UNICODE_ESCAPE_RE.sub(r'\\\\u', corrigido)
    corrigido_escapes = INVALID_ESCAPE_RE.sub(r'\\\\', corrigido_escapes)
    adicionar(corrigido_escapes)
    
    json_linha_unica = re.sub(r'(?<!\\)\r?\n', ' ', corrigido_escapes)
    json_linha_unica = re.sub(r'\s{2,}', ' ', json_linha_unica)
    adicionar(json_linha_unica)
    
    return candidatos


def parse_json_tolerante(resposta: str) -> Tuple[Optional[Dict], Dict[str, str]]:
    """Tenta interpretar a resposta como JSON, aplicando correcoes comuns."""
    debug = {
        "resposta_bruta": resposta or "",
        "json_extraido": ""
    }
    
    json_str = extrair_json_candidato(resposta)
    debug["json_extraido"] = json_str
    
    if not json_str:
        debug["erro_final"] = "Nenhum bloco JSON encontrado na resposta"
        return None, debug
    
    ultimo_erro = ""
    candidatos = gerar_candidatos_json(json_str)
    
    for idx, candidato in enumerate(candidatos, start=1):
        debug[f"json_candidato_{idx}"] = candidato
        try:
            return json.loads(candidato, strict=False), debug
        except json.JSONDecodeError as e:
            ultimo_erro = str(e)
            debug[f"erro_candidato_{idx}"] = ultimo_erro
    
    debug["erro_final"] = ultimo_erro or "Falha desconhecida ao interpretar JSON"
    return None, debug


def preparar_pasta_debug(pasta: Path) -> Tuple[Path, Path]:
    """Cria a pasta de debug desta execucao e retorna (dir, arquivo_log)."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    debug_dir = pasta / "_botsintese_debug" / timestamp
    debug_dir.mkdir(parents=True, exist_ok=True)
    return debug_dir, debug_dir / "execucao.log"


def salvar_debug_texto(debug_dir: Optional[Path], nome: str, conteudo: str) -> Optional[Path]:
    """Salva texto de debug em UTF-8."""
    if not debug_dir:
        return None
    
    caminho = debug_dir / nome
    caminho.write_text(conteudo or "", encoding='utf-8-sig')
    return caminho


def normalizar_nome(nome: str) -> str:
    """Normaliza nome de parte para evitar duplicatas por acento/caixa"""
    if not nome:
        return ""
    
    import unicodedata
    
    # Remove acentos
    nome_normalizado = unicodedata.normalize('NFKD', nome)
    nome_normalizado = ''.join(c for c in nome_normalizado if not unicodedata.combining(c))
    
    # Converte para maiúsculas
    nome_normalizado = nome_normalizado.upper().strip()
    
    # Remove variações comuns de sufixos empresariais
    sufixos = [' LTDA.', ' LTDA', ' S/A', ' S.A.', ' S.A', ' EPP', ' ME', ' EIRELI', ' SOCIEDADE SIMPLES']
    for sufixo in sufixos:
        nome_normalizado = nome_normalizado.replace(sufixo, '')
    
    # Remove espaços extras
    nome_normalizado = ' '.join(nome_normalizado.split())
    
    return nome_normalizado


def parse_data_brasileira(data_str: str) -> tuple:
    """Converte data dd/mm/aaaa para tupla ordenável (aaaa, mm, dd)"""
    if not data_str:
        return (0, 0, 0)
    
    try:
        # Tenta formato dd/mm/aaaa
        partes = data_str.strip().split('/')
        if len(partes) == 3:
            dia, mes, ano = partes
            return (int(ano), int(mes), int(dia))
    except:
        pass
    
    return (0, 0, 0)


def is_evento_relevante(descricao: str) -> bool:
    """Verifica se um evento do histórico é juridicamente relevante"""
    if not descricao:
        return False
    
    desc_lower = descricao.lower()
    
    # Termos que SEMPRE indicam relevância (allowlist - tem prioridade)
    termos_sempre_relevantes = [
        'despacho', 'decisão', 'sentença', 'acórdão',
        'citação', 'citado', 'cite-se', 'intimação', 'intimado',
        'penhora', 'penhorado', 'avaliação', 'leilão', 'hasta',
        'mandado', 'carta precatória', 'carta rogatória',
        'contestação', 'réplica', 'impugnação', 'embargos',
        'perícia', 'laudo', 'perito', 'audiência',
        'acordo', 'homologação', 'cumprimento',
        'recurso', 'apelação', 'agravo', 'tutela',
        'bloqueio', 'sisbajud', 'renajud', 'infojud',
        'dê-se vista', 'manifestação',
    ]
    
    for termo in termos_sempre_relevantes:
        if termo in desc_lower:
            return True
    
    # Lista de termos que indicam eventos IRRELEVANTES (ruído)
    termos_irrelevantes = [
        'assinado eletronicamente',
        'assinatura eletrônica',
        'documento assinado',
        'concluso para assinatura',
        'juntada automática',
        'certidão de publicação',
        'expediente forense',
        'não houve expediente',
        'feriado',
        'recesso',
        'portaria conjunta',
    ]
    
    for termo in termos_irrelevantes:
        if termo in desc_lower:
            return False
    
    return True


def categorizar_evento(descricao: str) -> str:
    """Categoriza evento como 'processual' ou 'fatico'"""
    if not descricao:
        return "processual"
    
    desc_lower = descricao.lower()
    
    # Termos que indicam eventos FÁTICOS (não processuais)
    termos_faticos = [
        'contrato', 'pagamento', 'pago', 'boleto', 'parcela',
        'protesto', 'negativação', 'serasa', 'spc', 'cadastro',
        'whatsapp', 'mensagem', 'email', 'notificação extrajudicial',
        'renegociação', 'acordo', 'tratamento', 'serviço',
        'emissão', 'vencimento', 'prestação',
    ]
    
    for termo in termos_faticos:
        if termo in desc_lower:
            return "fatico"
    
    return "processual"


def deduplicar_valores(valores: List[Dict]) -> List[Dict]:
    """Remove valores duplicados de forma mais inteligente"""
    if not valores:
        return []
    
    valores_unicos = []
    valores_vistos = {}  # valor_numerico -> item completo
    
    for v in valores:
        if not isinstance(v, dict):
            continue
        
        desc = v.get("descricao", "").strip()
        valor = v.get("valor", "").strip()
        
        if not desc or not valor:
            continue
        
        # Extrai valor numérico para comparação
        valor_num = re.sub(r'[^\d,.]', '', valor).replace('.', '').replace(',', '.')
        try:
            valor_float = float(valor_num) if valor_num else 0
        except:
            valor_float = 0
        
        # Chave composta: valor numérico + primeiras palavras da descrição
        palavras_chave = ' '.join(desc.lower().split()[:3])
        chave = f"{valor_float:.2f}|{palavras_chave}"
        
        if chave not in valores_vistos:
            valores_vistos[chave] = v
            valores_unicos.append(v)
    
    return valores_unicos


def mesclar_extracoes(extracoes: List[Dict]) -> Dict:
    """Mescla múltiplas extrações em uma única, via Python (sem IA) - v3.0"""
    
    resultado = {
        "partes": [],
        "objeto_acao": "",
        "resumo_fatos": "",
        "valores_relevantes": [],
        "pedidos": [],
        "pedidos_autor": [],
        "pedidos_reu": [],
        "pedidos_reconvencao": [],
        "preliminares": [],
        "decisoes": [],
        "teses_autor": [],
        "teses_reu": [],
        "documentos_importantes": [],
        "historico_processual": [],  # Atos do processo
        "historico_fatico": [],      # Linha do tempo dos fatos
        "status_atual": ""
    }
    
    # Sets para evitar duplicatas (usando nomes normalizados)
    partes_normalizadas = {}  # nome_normalizado -> dados originais
    pedidos_vistos = set()
    teses_autor_vistas = set()
    teses_reu_vistas = set()
    docs_vistos = set()
    historico_vistos = set()
    
    # Coleta todos os valores para deduplicar depois
    todos_valores = []
    
    for ext in extracoes:
        if not isinstance(ext, dict):
            continue
        
        # Objeto da ação - pega o primeiro não vazio
        if not resultado["objeto_acao"] and ext.get("objeto_acao"):
            resultado["objeto_acao"] = ext["objeto_acao"]
        
        # Resumo dos fatos - concatena se diferentes
        resumo = ext.get("resumo_fatos", "")
        if resumo and resumo not in resultado["resumo_fatos"]:
            if resultado["resumo_fatos"]:
                resultado["resumo_fatos"] += "\n\n" + resumo
            else:
                resultado["resumo_fatos"] = resumo
        
        # Partes - normaliza e evita duplicatas
        for p in ext.get("partes", []):
            if isinstance(p, dict):
                nome_original = p.get("nome", "").strip()
                if not nome_original or nome_original.upper() == "NONE":
                    continue
                
                nome_norm = normalizar_nome(nome_original)
                if nome_norm and nome_norm not in partes_normalizadas:
                    partes_normalizadas[nome_norm] = p
        
        # Valores - coleta todos para deduplicar depois
        for v in ext.get("valores_relevantes", []):
            if isinstance(v, dict):
                todos_valores.append(v)
        
        # Pedidos
        for p in ext.get("pedidos", []):
            if p and isinstance(p, str):
                p_lower = p.lower().strip()[:50]
                if p_lower not in pedidos_vistos:
                    pedidos_vistos.add(p_lower)
                    resultado["pedidos"].append(p)
        
        # Decisões
        for d in ext.get("decisoes", []):
            if isinstance(d, dict):
                resultado["decisoes"].append(d)
        
        # Teses do autor
        for t in ext.get("teses_autor", []):
            if t and isinstance(t, str):
                t_lower = t.lower().strip()[:50]
                if t_lower not in teses_autor_vistas:
                    teses_autor_vistas.add(t_lower)
                    resultado["teses_autor"].append(t)
        
        # Teses do réu
        for t in ext.get("teses_reu", []):
            if t and isinstance(t, str):
                t_lower = t.lower().strip()[:50]
                if t_lower not in teses_reu_vistas:
                    teses_reu_vistas.add(t_lower)
                    resultado["teses_reu"].append(t)
        
        # Documentos importantes
        for d in ext.get("documentos_importantes", []):
            if isinstance(d, dict):
                tipo = d.get("tipo", "").lower().strip()
                if tipo and tipo not in docs_vistos:
                    docs_vistos.add(tipo)
                    resultado["documentos_importantes"].append(d)
        
        # Histórico - filtra e categoriza
        for h in ext.get("historico_detalhado", []):
            if isinstance(h, dict):
                data = h.get("data", "")
                desc = h.get("descricao", h.get("evento", ""))
                
                # Pula eventos irrelevantes
                if not is_evento_relevante(desc):
                    continue
                
                chave = f"{data}|{desc[:80]}".lower()
                if chave not in historico_vistos:
                    historico_vistos.add(chave)
                    
                    # Categoriza e adiciona na lista correta
                    categoria = categorizar_evento(desc)
                    if categoria == "fatico":
                        resultado["historico_fatico"].append(h)
                    else:
                        resultado["historico_processual"].append(h)
        
        # Status - pega o último não vazio
        if ext.get("status_atual"):
            resultado["status_atual"] = ext["status_atual"]
    
    # Finaliza partes (sem duplicatas)
    resultado["partes"] = list(partes_normalizadas.values())
    
    # Deduplica valores de forma inteligente
    resultado["valores_relevantes"] = deduplicar_valores(todos_valores)
    
    # Ordena históricos por data (cronológico)
    for campo in ["historico_processual", "historico_fatico"]:
        try:
            resultado[campo].sort(key=lambda x: parse_data_brasileira(x.get("data", "")))
        except:
            pass
    
    # Mantém compatibilidade: histórico completo ordenado
    resultado["historico_detalhado"] = []
    resultado["historico_detalhado"].extend(resultado["historico_processual"])
    resultado["historico_detalhado"].extend(resultado["historico_fatico"])
    try:
        resultado["historico_detalhado"].sort(key=lambda x: parse_data_brasileira(x.get("data", "")))
    except:
        pass
    
    return resultado


# ============================================================================
# PROVEDORES DE LLM
# ============================================================================

def chamar_ollama(prompt: str, config: Config, timeout: int = 180) -> str:
    """Chama modelo local via Ollama"""
    try:
        r = requests.post(
            f"{config.ollama_host}/api/generate",
            json={
                "model": config.modelo_local,
                "prompt": prompt,
                "stream": False,
                "options": {"temperature": 0.2, "num_predict": 2000}
            },
            timeout=timeout
        )
        if r.status_code == 200:
            return r.json().get('response', '').strip()
    except Exception as e:
        print(f"Erro Ollama: {e}")
    return ""


# Controle de rate limiting para Google API gratuita
_google_last_request = 0
_google_request_count = 0
_google_minute_start = 0

def chamar_google(prompt: str, config: Config, retry_count: int = 0) -> str:
    """Chama Google Gemini API com rate limiting para plano gratuito"""
    global _google_last_request, _google_request_count, _google_minute_start
    
    if not config.api_google:
        raise ValueError("API key do Google não configurada")
    
    # Rate limiting: máximo 15 requisições por minuto no plano gratuito
    current_time = time.time()
    
    # Reset contador se passou 1 minuto
    if current_time - _google_minute_start > 60:
        _google_request_count = 0
        _google_minute_start = current_time
    
    # Se atingiu limite, espera
    if _google_request_count >= 14:  # 14 para margem de segurança
        wait_time = 60 - (current_time - _google_minute_start) + 1
        if wait_time > 0:
            print(f"    ⏳ Rate limit: aguardando {wait_time:.0f}s...")
            time.sleep(wait_time)
            _google_request_count = 0
            _google_minute_start = time.time()
    
    # Pausa mínima entre requisições (4 segundos = ~15/min)
    time_since_last = current_time - _google_last_request
    if time_since_last < 4 and _google_last_request > 0:
        time.sleep(4 - time_since_last)
    
    # Modelo atualizado para 2026 (gemini-1.5-flash foi desativado em Set/2025)
    modelo = "gemini-2.5-flash"
    
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{modelo}:generateContent?key={config.api_google}"
    
    try:
        _google_last_request = time.time()
        _google_request_count += 1
        
        r = requests.post(
            url,
            json={
                "contents": [{"parts": [{"text": prompt}]}],
                "generationConfig": {
                    "temperature": 0.2,
                    "maxOutputTokens": 65536,
                    "responseMimeType": "application/json"  # Força JSON válido
                }
            },
            timeout=180  # Timeout maior para textos grandes
        )
        
        if r.status_code == 200:
            data = r.json()
            if 'candidates' in data and len(data['candidates']) > 0:
                return data['candidates'][0]['content']['parts'][0]['text']
            else:
                print(f"    ⚠️ Resposta vazia do Gemini")
                return ""
        
        elif r.status_code == 429:  # Rate limit exceeded
            if retry_count < 3:
                print(f"    ⚠️ Rate limit atingido, aguardando 60s... (tentativa {retry_count + 1}/3)")
                time.sleep(60)
                _google_request_count = 0
                _google_minute_start = time.time()
                return chamar_google(prompt, config, retry_count + 1)
            else:
                print(f"    ❌ Rate limit persistente após 3 tentativas")
                return ""
        
        elif r.status_code == 400:
            error_msg = r.json().get('error', {}).get('message', r.text)
            print(f"    ⚠️ Erro 400: {error_msg[:100]}")
            # Se o prompt for muito grande, pode ser erro de tamanho
            if "too long" in error_msg.lower() or "token" in error_msg.lower():
                print(f"    💡 Texto pode ser muito grande, tente dividir o processo")
            return ""
        
        else:
            print(f"Erro Google API: {r.status_code} - {r.text[:200]}")
            
    except requests.exceptions.Timeout:
        print(f"    ⚠️ Timeout na requisição (180s)")
        if retry_count < 2:
            return chamar_google(prompt, config, retry_count + 1)
    except Exception as e:
        print(f"Erro Google: {e}")
    
    return ""


def chamar_anthropic(prompt: str, config: Config) -> str:
    """Chama Anthropic Claude API"""
    if not config.api_anthropic:
        raise ValueError("API key da Anthropic não configurada")
    
    try:
        r = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "x-api-key": config.api_anthropic,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json"
            },
            json={
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 4000,
                "messages": [{"role": "user", "content": prompt}]
            },
            timeout=120
        )
        if r.status_code == 200:
            data = r.json()
            return data['content'][0]['text']
        else:
            print(f"Erro Anthropic API: {r.status_code} - {r.text}")
    except Exception as e:
        print(f"Erro Anthropic: {e}")
    return ""


def chamar_openai(prompt: str, config: Config) -> str:
    """Chama OpenAI GPT API"""
    if not config.api_openai:
        raise ValueError("API key da OpenAI não configurada")
    
    try:
        r = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {config.api_openai}",
                "Content-Type": "application/json"
            },
            json={
                "model": "gpt-4o",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.2,
                "max_tokens": 4000
            },
            timeout=120
        )
        if r.status_code == 200:
            data = r.json()
            return data['choices'][0]['message']['content']
        else:
            print(f"Erro OpenAI API: {r.status_code} - {r.text}")
    except Exception as e:
        print(f"Erro OpenAI: {e}")
    return ""


def chamar_xai(prompt: str, config: Config) -> str:
    """Chama xAI Grok API"""
    if not config.api_xai:
        raise ValueError("API key da xAI não configurada")
    
    try:
        r = requests.post(
            "https://api.x.ai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {config.api_xai}",
                "Content-Type": "application/json"
            },
            json={
                "model": "grok-beta",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.2,
                "max_tokens": 4000
            },
            timeout=120
        )
        if r.status_code == 200:
            data = r.json()
            return data['choices'][0]['message']['content']
        else:
            print(f"Erro xAI API: {r.status_code} - {r.text}")
    except Exception as e:
        print(f"Erro xAI: {e}")
    return ""


def chamar_llm(prompt: str, modo: str, config: Config) -> str:
    """Chama o LLM apropriado baseado no modo selecionado"""
    if modo == "local":
        return chamar_ollama(prompt, config)
    elif modo == "google":
        return chamar_google(prompt, config)
    elif modo == "anthropic":
        return chamar_anthropic(prompt, config)
    elif modo == "openai":
        return chamar_openai(prompt, config)
    elif modo == "xai":
        return chamar_xai(prompt, config)
    else:
        raise ValueError(f"Modo desconhecido: {modo}")


# ============================================================================
# PROCESSAMENTO PRINCIPAL
# ============================================================================

def processar_processo(pasta: Path, modo: str, config: Config, callback=None) -> Dict:
    """Processa todos os PDFs de uma pasta"""
    debug_dir, log_file = preparar_pasta_debug(pasta)
    
    def log(msg):
        try:
            print(msg)
        except UnicodeEncodeError:
            encoding = sys.stdout.encoding or 'utf-8'
            msg_segura = msg.encode(encoding, errors='replace').decode(encoding, errors='replace')
            print(msg_segura)
        with open(log_file, 'a', encoding='utf-8-sig') as f:
            f.write(msg + "\n")
        if callback:
            callback(msg)
    
    resultado = {
        'dados': DadosProcesso(),
        'extracao': {},
        'tempo': 0,
        'modo': modo
    }
    
    inicio = time.time()
    log(f"🗂️ Log detalhado: {log_file}")
    
    # Encontra PDFs (incluindo subpastas)
    pdfs = list(pasta.glob("*.pdf"))
    pdfs += list(pasta.glob("**/*.pdf"))  # Subpastas também
    pdfs = list(set(pdfs))  # Remove duplicatas de caminho
    
    if not pdfs:
        log("❌ Nenhum PDF encontrado!")
        return resultado
    
    log(f"📄 Encontrados {len(pdfs)} PDFs")
    
    # Identifica arquivos importantes (prefixo ou subpasta)
    prefixos_importantes = ("IMPORTANTE_", "PRINCIPAL_", "DESTAQUE_")
    pasta_importantes = "importantes"
    
    arquivos_importantes = []
    arquivos_normais = []
    
    for pdf in pdfs:
        eh_importante = (
            any(pdf.name.startswith(p) for p in prefixos_importantes) or
            pasta_importantes in str(pdf.parent).lower()
        )
        if eh_importante:
            arquivos_importantes.append(pdf)
            log(f"  ⭐ {pdf.name} (IMPORTANTE)")
        else:
            arquivos_normais.append(pdf)
            log(f"  📄 {pdf.name}")
    
    # Extrai texto de todos os PDFs e deduplica por conteúdo
    textos_unicos = {}  # hash -> (nome, texto, importante)
    total_paginas = 0
    
    # Processa importantes primeiro
    for pdf in arquivos_importantes + arquivos_normais:
        texto, num_pag = extrair_texto_pdf(pdf)
        if texto.strip():
            # Gera hash do conteúdo para detectar duplicatas
            texto_hash = hashlib.md5(texto[:10000].encode()).hexdigest()
            eh_importante = pdf in arquivos_importantes
            
            if texto_hash not in textos_unicos:
                textos_unicos[texto_hash] = (pdf.name, texto, eh_importante)
                total_paginas += num_pag
            else:
                # Se já existe, mas este é importante e o outro não, substitui
                nome_existente, texto_existente, importante_existente = textos_unicos[texto_hash]
                if eh_importante and not importante_existente:
                    textos_unicos[texto_hash] = (pdf.name, texto, True)
                    log(f"    ⚠️ Conteúdo igual a '{nome_existente}', mas este é IMPORTANTE - usando este")
                else:
                    log(f"    ⚠️ Conteúdo duplicado de '{nome_existente}' - ignorando")
        else:
            log(f"    ⚠️ {pdf.name}: Sem texto extraível (verifique o OCR)")
    
    log(f"📊 Total: {total_paginas} páginas ({len(textos_unicos)} documentos únicos)")
    
    # Conta importantes
    num_importantes = sum(1 for _, _, imp in textos_unicos.values() if imp)
    if num_importantes > 0:
        log(f"⭐ {num_importantes} documentos marcados como importantes")
    
    if not textos_unicos:
        log("❌ Nenhum texto extraído! Verifique o OCR.")
        return resultado
    
    # Junta todos os textos únicos (importantes primeiro)
    textos_ordenados = sorted(textos_unicos.values(), key=lambda x: (not x[2], x[0]))  # Importantes primeiro
    texto_completo = "\n\n".join([t[1] for t in textos_ordenados])
    
    # Detecta sistema e extrai dados estruturados
    log("\n🔍 Detectando sistema processual...")
    sistema = detectar_sistema(texto_completo)
    log(f"  Sistema identificado: {sistema.upper()}")
    
    if sistema == "pje":
        dados = extrair_dados_pje(texto_completo)
    elif sistema == "eproc":
        dados = extrair_dados_eproc(texto_completo)
    elif sistema == "saj":
        dados = extrair_dados_saj(texto_completo)
    else:
        dados = extrair_dados_generico(texto_completo)
    
    # Deduplica eventos também
    eventos_unicos = []
    eventos_vistos = set()
    for e in dados.eventos:
        chave = f"{e.data}|{e.tipo}|{e.descricao}"
        if chave not in eventos_vistos:
            eventos_vistos.add(chave)
            eventos_unicos.append(e)
    dados.eventos = eventos_unicos
    
    log(f"  Processo: {dados.numero or 'não identificado'}")
    log(f"  Eventos encontrados: {len(dados.eventos)}")
    
    # Divide em chunks para análise (tamanho depende do modo)
    chunks = dividir_em_chunks(texto_completo, config, modo)
    log(f"\n📝 Dividido em {len(chunks)} partes para análise")
    
    # Se for cloud e tiver poucos chunks, pode processar tudo de uma vez
    if modo in ["google", "anthropic", "openai", "xai"] and len(chunks) <= 3:
        log(f"   💡 Contexto grande disponível - processamento otimizado")
    
    # Extrai informações de cada chunk
    log(f"\n🤖 Processando com {modo.upper()}...")
    extracoes = []
    
    for i, chunk in enumerate(chunks):
        log(f"  Parte {i+1}/{len(chunks)}...")
        
        prompt = PROMPT_EXTRACAO.format(texto=chunk)
        resposta = chamar_llm(prompt, modo, config)
        
        if not resposta:
            log(f"    ⚠️ Modelo retornou resposta vazia")
            continue
        
        extracao, debug_json = parse_json_tolerante(resposta)
        if extracao:
            extracoes.append(extracao)
            if "erro_candidato_1" in debug_json:
                log(f"    ✅ Extraído após correção automática")
            else:
                log(f"    ✅ Extraído com sucesso")
            continue
        
        erro_inicial = debug_json.get("erro_candidato_1")
        if erro_inicial:
            log(f"    ⚠️ Resposta não é JSON válido: {erro_inicial[:80]}")
        else:
            log(f"    ⚠️ Nenhum JSON encontrado na resposta")
            log(f"    📝 Início da resposta: {resposta[:200]}...")
        
        base_nome = f"parte_{i+1:02d}"
        bruto_path = salvar_debug_texto(debug_dir, f"{base_nome}_resposta_bruta.txt", debug_json.get("resposta_bruta", ""))
        salvar_debug_texto(debug_dir, f"{base_nome}_json_extraido.txt", debug_json.get("json_extraido", ""))
        for chave, conteudo in debug_json.items():
            if chave.startswith("json_candidato_"):
                salvar_debug_texto(debug_dir, f"{base_nome}_{chave}.txt", conteudo)
        
        log(f"    ❌ Falha definitiva no parse: {debug_json.get('erro_final', 'erro desconhecido')[:80]}")
        if bruto_path:
            log(f"    🧪 Resposta bruta salva em: {bruto_path}")
        continue
        
        if resposta:
            # Tenta extrair JSON da resposta
            try:
                # Remove code blocks markdown se existirem
                resposta_limpa = resposta
                if '```json' in resposta_limpa:
                    resposta_limpa = re.sub(r'```json\s*', '', resposta_limpa)
                    resposta_limpa = re.sub(r'```\s*$', '', resposta_limpa)
                elif '```' in resposta_limpa:
                    resposta_limpa = re.sub(r'```\s*', '', resposta_limpa)
                
                # Procura por JSON na resposta
                json_match = re.search(r'\{[\s\S]*\}', resposta_limpa)
                if json_match:
                    json_str = json_match.group()
                    extracao = json.loads(json_str)
                    extracoes.append(extracao)
                    log(f"    ✅ Extraído com sucesso")
                else:
                    log(f"    ⚠️ Nenhum JSON encontrado na resposta")
                    # Log primeiros 200 chars para debug
                    log(f"    📝 Início da resposta: {resposta[:200]}...")
            except json.JSONDecodeError as e:
                log(f"    ⚠️ Resposta não é JSON válido: {str(e)[:50]}")
                # Tenta corrigir erros comuns de JSON
                try:
                    json_str = resposta
                    # Remove texto antes/depois do JSON
                    inicio_json = json_str.find('{')
                    fim_json = json_str.rfind('}')
                    if inicio_json >= 0 and fim_json > inicio_json:
                        json_str = json_str[inicio_json:fim_json+1]
                    
                    # Correções comuns de JSON malformado
                    # 1. Remove vírgulas antes de } ou ]
                    json_str = re.sub(r',\s*}', '}', json_str)
                    json_str = re.sub(r',\s*]', ']', json_str)
                    # 2. Adiciona vírgulas faltando entre } e "
                    json_str = re.sub(r'}\s*"', '}, "', json_str)
                    json_str = re.sub(r']\s*"', '], "', json_str)
                    # 3. Corrige aspas não escapadas dentro de strings (mais complexo)
                    # 4. Remove quebras de linha dentro de strings
                    json_str = re.sub(r'(?<!\\)\n', ' ', json_str)
                    
                    extracao = json.loads(json_str)
                    extracoes.append(extracao)
                    log(f"    ✅ Extraído após correção automática")
                except Exception as e2:
                    log(f"    ❌ Falha definitiva no parse: {str(e2)[:30]}")
    
    # Consolida extrações via Python (mais rápido e confiável que IA)
    if len(extracoes) > 1:
        log("\n📋 Consolidando informações...")
        resultado['extracao'] = mesclar_extracoes(extracoes)
        log("    ✅ Consolidação OK")
    elif extracoes:
        resultado['extracao'] = extracoes[0]
    elif not extracoes:
        log("    ⚠️ Nenhuma extração bem-sucedida - relatório pode ficar incompleto")
    
    resultado['dados'] = dados
    resultado['tempo'] = time.time() - inicio
    
    log(f"\n✅ Concluído em {resultado['tempo']:.1f} segundos")
    
    return resultado


# ============================================================================
# GERAÇÃO DE RELATÓRIOS
# ============================================================================

def gerar_markdown(resultado: Dict, pasta: Path) -> str:
    """Gera relatório em Markdown"""
    dados = resultado['dados']
    extracao = resultado.get('extracao', {})
    
    md = []
    
    # Cabeçalho
    md.append("# Síntese Processual")
    md.append(f"**Processo:** {dados.numero or 'Não identificado'}")
    md.append(f"**Gerado em:** {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
    md.append(f"**Modo:** {resultado.get('modo', 'N/D').upper()}")
    md.append(f"**Tempo de processamento:** {resultado.get('tempo', 0):.1f} segundos")
    md.append("")
    md.append("---")
    md.append("")
    
    # Dados Gerais (regex + fallback para extração do LLM)
    md.append("## Dados Gerais")
    md.append("")
    classe = dados.classe or extracao.get('classe_processual', '')
    if classe:
        md.append(f"- **Classe:** {classe}")
    vara = dados.vara or extracao.get('vara', '')
    if vara:
        md.append(f"- **Vara:** {vara}")
    comarca = dados.comarca or extracao.get('comarca', '')
    if comarca:
        md.append(f"- **Comarca:** {comarca}")
    if dados.valor_causa:
        md.append(f"- **Valor da causa:** {dados.valor_causa}")
    distribuicao = dados.data_distribuicao or extracao.get('data_distribuicao', '')
    if distribuicao:
        md.append(f"- **Distribuição:** {distribuicao}")
    assunto = dados.assunto or extracao.get('assunto', '')
    if assunto:
        md.append(f"- **Assunto:** {assunto}")
    md.append("")
    
    # Partes
    md.append("## Partes")
    md.append("")
    
    partes = extracao.get('partes_consolidadas') or extracao.get('partes') or dados.partes
    if partes:
        md.append("| Polo | Nome |")
        md.append("|------|------|")
        for p in partes:
            if isinstance(p, dict):
                nome = p.get('nome', 'N/D')
                if nome and nome != 'None' and nome != 'null':
                    md.append(f"| {p.get('polo', 'N/D')} | {nome} |")
        md.append("")
    
    # Objeto da Ação
    objeto = extracao.get('objeto_acao', '')
    if objeto:
        md.append("## Objeto da Ação")
        md.append("")
        md.append(objeto)
        md.append("")
    
    # Resumo dos Fatos (com parágrafos)
    resumo = extracao.get('resumo_fatos', '')
    
    if resumo:
        md.append("## Resumo dos Fatos")
        md.append("")
        # Garante que há quebras de parágrafo
        resumo_formatado = resumo.replace('\\n\\n', '\n\n').replace('\\n', '\n')
        # Se não tem parágrafos, tenta dividir em sentenças longas
        if '\n\n' not in resumo_formatado and len(resumo_formatado) > 500:
            # Divide em parágrafos a cada ~300 caracteres no ponto final
            partes = []
            atual = ""
            for frase in resumo_formatado.split('. '):
                atual += frase + '. '
                if len(atual) > 300:
                    partes.append(atual.strip())
                    atual = ""
            if atual.strip():
                partes.append(atual.strip())
            resumo_formatado = '\n\n'.join(partes)
        md.append(resumo_formatado)
        md.append("")
    
    # Documentos Importantes (NOVA SEÇÃO)
    docs_importantes = extracao.get('documentos_importantes', [])
    if docs_importantes:
        md.append("## Documentos Importantes")
        md.append("")
        for i, doc in enumerate(docs_importantes, 1):
            if isinstance(doc, dict):
                tipo = doc.get('tipo', 'Documento')
                data = doc.get('data', '')
                parte = doc.get('parte', '')
                resumo_doc = doc.get('resumo', '')
                
                titulo = f"### {i}. {tipo}"
                if data:
                    titulo += f" ({data})"
                md.append(titulo)
                if parte:
                    md.append(f"**Apresentado por:** {parte}")
                md.append("")
                if resumo_doc:
                    md.append(resumo_doc)
                md.append("")
        md.append("---")
        md.append("")
    
    # Histórico Processual (atos do processo)
    historico_proc = extracao.get('historico_processual', [])
    historico_geral = extracao.get('historico_resumido') or extracao.get('historico_detalhado', [])
    
    # Se tiver histórico processual separado, usa ele
    if historico_proc:
        md.append("## Histórico Processual")
        md.append("")
        md.append("| Data | Descrição |")
        md.append("|------|-----------|")
        for h in historico_proc:
            if isinstance(h, dict):
                data = h.get('data', 'N/D')
                desc = h.get('descricao', h.get('evento', 'N/D'))
                md.append(f"| {data} | {desc} |")
        md.append("")
    elif historico_geral:
        md.append("## Histórico Processual")
        md.append("")
        md.append("| Data | Descrição |")
        md.append("|------|-----------|")
        for h in historico_geral:
            if isinstance(h, dict):
                data = h.get('data', 'N/D')
                desc = h.get('descricao', h.get('evento', 'N/D'))
                md.append(f"| {data} | {desc} |")
        md.append("")
    elif dados.eventos:
        md.append("## Histórico Processual")
        md.append("")
        md.append("| Data | Tipo | Descrição |")
        md.append("|------|------|-----------|")
        for e in dados.eventos[:30]:
            md.append(f"| {e.data} | {e.tipo} | {e.descricao[:60]} |")
        md.append("")
    
    # Linha do Tempo Fática (se houver)
    historico_fatico = extracao.get('historico_fatico', [])
    if historico_fatico:
        md.append("## Linha do Tempo dos Fatos")
        md.append("")
        md.append("| Data | Descrição |")
        md.append("|------|-----------|")
        for h in historico_fatico:
            if isinstance(h, dict):
                data = h.get('data', 'N/D')
                desc = h.get('descricao', h.get('evento', 'N/D'))
                md.append(f"| {data} | {desc} |")
        md.append("")
    
    # Valores
    valores = extracao.get('valores_relevantes') or extracao.get('valores_consolidados') or extracao.get('valores', [])
    if valores:
        md.append("## Valores Identificados")
        md.append("")
        for v in valores:
            if isinstance(v, dict):
                md.append(f"- **{v.get('descricao', 'N/D')}:** {v.get('valor', 'N/D')}")
        md.append("")
    
    # Teses das Partes
    teses_autor = extracao.get('teses_autor', [])
    teses_reu = extracao.get('teses_reu', [])
    
    if teses_autor or teses_reu:
        md.append("## Teses das Partes")
        md.append("")
        
        if teses_autor:
            md.append("**Autor:**")
            for t in teses_autor:
                md.append(f"- {t}")
            md.append("")
        
        if teses_reu:
            md.append("**Réu:**")
            for t in teses_reu:
                md.append(f"- {t}")
            md.append("")
    
    # Decisões
    decisoes = extracao.get('decisoes_importantes') or extracao.get('decisoes', [])
    if decisoes:
        md.append("## Decisões")
        md.append("")
        for d in decisoes:
            if isinstance(d, dict):
                data = d.get('data', 'N/D')
                tipo = d.get('tipo', 'N/D')
                conteudo = d.get('conteudo', 'N/D')
                if conteudo and conteudo != 'None' and conteudo != 'null':
                    md.append(f"- **{data} - {tipo}:** {conteudo}")
        md.append("")
    
    # Status
    status = extracao.get('status_atual', '')
    if status:
        md.append("## Status Atual")
        md.append("")
        md.append(status)
        md.append("")
    
    # Rodapé
    md.append("---")
    md.append("")
    md.append("*Documento gerado automaticamente pelo BotSíntese v3.0*")
    md.append("*Este é um resumo factual. Não contém análises ou recomendações jurídicas.*")
    
    return "\n".join(md)


def gerar_docx(resultado: Dict, pasta: Path) -> Document:
    """Gera relatório em Word"""
    doc = Document()
    dados = resultado['dados']
    extracao = resultado.get('extracao', {})
    
    # Título
    titulo = doc.add_heading('Síntese Processual', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadados
    doc.add_paragraph(f"Processo: {dados.numero or 'Não identificado'}")
    doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
    doc.add_paragraph(f"Modo: {resultado.get('modo', 'N/D').upper()}")
    
    doc.add_paragraph("─" * 50)
    
    # Dados Gerais
    doc.add_heading('Dados Gerais', level=1)
    if dados.classe:
        doc.add_paragraph(f"Classe: {dados.classe}")
    if dados.vara:
        doc.add_paragraph(f"Vara: {dados.vara}")
    if dados.valor_causa:
        doc.add_paragraph(f"Valor da causa: {dados.valor_causa}")
    if dados.data_distribuicao:
        doc.add_paragraph(f"Distribuição: {dados.data_distribuicao}")
    
    # Partes
    doc.add_heading('Partes', level=1)
    partes = extracao.get('partes_consolidadas') or extracao.get('partes') or dados.partes
    if partes:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Polo'
        hdr[1].text = 'Nome'
        for p in partes:
            if isinstance(p, dict):
                nome = p.get('nome', 'N/D')
                if nome and nome != 'None' and nome != 'null':
                    row = table.add_row().cells
                    row[0].text = str(p.get('polo', 'N/D') or 'N/D')
                    row[1].text = str(nome or 'N/D')
    
    # Objeto
    objeto = extracao.get('objeto_acao', '')
    if objeto:
        doc.add_heading('Objeto da Ação', level=1)
        doc.add_paragraph(objeto)
    
    # Resumo dos Fatos (com parágrafos)
    resumo = extracao.get('resumo_fatos', '')
    if resumo:
        doc.add_heading('Resumo dos Fatos', level=1)
        # Formata parágrafos
        resumo_formatado = resumo.replace('\\n\\n', '\n\n').replace('\\n', '\n')
        for paragrafo in resumo_formatado.split('\n\n'):
            if paragrafo.strip():
                doc.add_paragraph(paragrafo.strip())
    
    # Documentos Importantes
    docs_importantes = extracao.get('documentos_importantes', [])
    if docs_importantes:
        doc.add_heading('Documentos Importantes', level=1)
        for i, doc_imp in enumerate(docs_importantes, 1):
            if isinstance(doc_imp, dict):
                tipo = doc_imp.get('tipo', 'Documento')
                data = doc_imp.get('data', '')
                parte = doc_imp.get('parte', '')
                resumo_doc = doc_imp.get('resumo', '')
                
                titulo_doc = f"{i}. {tipo}"
                if data:
                    titulo_doc += f" ({data})"
                doc.add_heading(titulo_doc, level=2)
                if parte:
                    doc.add_paragraph(f"Apresentado por: {parte}")
                if resumo_doc:
                    doc.add_paragraph(resumo_doc)
    
    # Histórico Processual
    historico_proc = extracao.get('historico_processual', [])
    historico_geral = extracao.get('historico_resumido') or extracao.get('historico_detalhado', [])
    
    if historico_proc:
        doc.add_heading('Histórico Processual', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Data'
        hdr[1].text = 'Descrição'
        for h in historico_proc:
            if isinstance(h, dict):
                row = table.add_row().cells
                row[0].text = str(h.get('data', 'N/D') or 'N/D')
                row[1].text = str(h.get('descricao', h.get('evento', 'N/D')) or 'N/D')
    elif historico_geral:
        doc.add_heading('Histórico Processual', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Data'
        hdr[1].text = 'Descrição'
        for h in historico_geral:
            if isinstance(h, dict):
                row = table.add_row().cells
                row[0].text = str(h.get('data', 'N/D') or 'N/D')
                row[1].text = str(h.get('descricao', h.get('evento', 'N/D')) or 'N/D')
    elif dados.eventos:
        doc.add_heading('Histórico Processual', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Data'
        hdr[1].text = 'Tipo'
        hdr[2].text = 'Descrição'
        for e in dados.eventos[:30]:
            row = table.add_row().cells
            row[0].text = str(e.data or 'N/D')
            row[1].text = str(e.tipo or 'N/D')
            row[2].text = str(e.descricao[:50] if e.descricao else 'N/D')
    
    # Linha do Tempo Fática (se houver)
    historico_fatico = extracao.get('historico_fatico', [])
    if historico_fatico:
        doc.add_heading('Linha do Tempo dos Fatos', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Data'
        hdr[1].text = 'Descrição'
        for h in historico_fatico:
            if isinstance(h, dict):
                row = table.add_row().cells
                row[0].text = str(h.get('data', 'N/D') or 'N/D')
                row[1].text = str(h.get('descricao', h.get('evento', 'N/D')) or 'N/D')
    
    # Valores
    valores = extracao.get('valores_relevantes') or extracao.get('valores_consolidados', [])
    if valores:
        doc.add_heading('Valores Identificados', level=1)
        for v in valores:
            if isinstance(v, dict):
                doc.add_paragraph(f"• {v.get('descricao', 'N/D')}: {v.get('valor', 'N/D')}")
    
    # Teses
    teses_autor = extracao.get('teses_autor', [])
    teses_reu = extracao.get('teses_reu', [])
    if teses_autor or teses_reu:
        doc.add_heading('Teses das Partes', level=1)
        if teses_autor:
            p = doc.add_paragraph()
            p.add_run("Autor:").bold = True
            for t in teses_autor:
                doc.add_paragraph(f"• {t}")
        if teses_reu:
            p = doc.add_paragraph()
            p.add_run("Réu:").bold = True
            for t in teses_reu:
                doc.add_paragraph(f"• {t}")
    
    # Rodapé
    doc.add_paragraph()
    doc.add_paragraph("─" * 50)
    p = doc.add_paragraph()
    p.add_run("Documento gerado automaticamente pelo BotSíntese v3.0").italic = True
    
    return doc


# ============================================================================
# INTERFACE GRÁFICA
# ============================================================================

class BotSinteseGUI:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("BotSíntese v3.0 - Síntese Processual")
        self.root.geometry("700x550")
        self.root.resizable(True, True)
        
        self.pasta_script = Path(__file__).parent
        self.config = carregar_config(self.pasta_script)
        self.pasta_selecionada = None
        self.processando = False
        
        self.criar_widgets()
    
    def criar_widgets(self):
        # Frame principal
        main = tk.Frame(self.root, padx=20, pady=15)
        main.pack(fill=tk.BOTH, expand=True)
        
        # Título
        tk.Label(main, text="BotSíntese v3.0", font=("Arial", 22, "bold")).pack()
        tk.Label(main, text="Síntese Processual Automatizada", font=("Arial", 11)).pack()
        
        # Seleção de pasta
        frame_pasta = tk.Frame(main)
        frame_pasta.pack(pady=15, fill=tk.X)
        
        self.btn_pasta = tk.Button(
            frame_pasta, text="📁 Selecionar Pasta do Processo",
            command=self.selecionar_pasta, font=("Arial", 11), padx=15, pady=8
        )
        self.btn_pasta.pack()
        
        self.lbl_pasta = tk.Label(frame_pasta, text="Nenhuma pasta selecionada", fg="gray")
        self.lbl_pasta.pack(pady=5)
        
        # Seleção de modo
        frame_modo = tk.LabelFrame(main, text="Modo de Processamento", padx=15, pady=10)
        frame_modo.pack(pady=10, fill=tk.X)
        
        self.modo_var = tk.StringVar(value=self.config.modo_padrao)
        
        modos = [
            ("google", "☁️ Google Gemini - GRATUITO (recomendado)"),
            ("local", "🖥️ Local (Ollama) - Gratuito, mais lento"),
            ("anthropic", "☁️ Anthropic Claude - ~R$ 1-3/processo"),
            ("openai", "☁️ OpenAI GPT-4o - ~R$ 2-5/processo"),
            ("xai", "☁️ xAI Grok - ~R$ 1-3/processo"),
        ]
        
        for valor, texto in modos:
            rb = tk.Radiobutton(
                frame_modo, text=texto, variable=self.modo_var, value=valor,
                font=("Arial", 10), anchor="w"
            )
            rb.pack(fill=tk.X)
        
        # Botão configurar APIs
        tk.Button(
            frame_modo, text="⚙️ Configurar APIs",
            command=self.abrir_config, font=("Arial", 9)
        ).pack(pady=5)
        
        # Botão processar
        self.btn_processar = tk.Button(
            main, text="▶️ Gerar Síntese", command=self.iniciar_processamento,
            font=("Arial", 13, "bold"), padx=25, pady=12, state=tk.DISABLED
        )
        self.btn_processar.pack(pady=10)
        
        # Log
        frame_log = tk.Frame(main)
        frame_log.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.log_text = tk.Text(frame_log, height=10, font=("Consolas", 9))
        self.log_text.pack(fill=tk.BOTH, expand=True, side=tk.LEFT)
        
        scrollbar = tk.Scrollbar(frame_log, command=self.log_text.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.log_text.config(yscrollcommand=scrollbar.set)
        
        # Status
        self.lbl_status = tk.Label(main, text="Pronto", font=("Arial", 10))
        self.lbl_status.pack()
    
    def log(self, msg):
        self.log_text.insert(tk.END, msg + "\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()
    
    def selecionar_pasta(self):
        pasta = filedialog.askdirectory(title="Selecione a pasta com os PDFs")
        if pasta:
            self.pasta_selecionada = Path(pasta)
            self.lbl_pasta.config(text=str(self.pasta_selecionada), fg="black")
            self.btn_processar.config(state=tk.NORMAL)
            
            pdfs = list(self.pasta_selecionada.glob("*.pdf"))
            self.log(f"Pasta: {pasta}")
            self.log(f"PDFs encontrados: {len(pdfs)}")
    
    def abrir_config(self):
        """Abre janela de configuração de APIs"""
        win = tk.Toplevel(self.root)
        win.title("Configurar APIs")
        win.geometry("500x400")
        win.transient(self.root)
        
        frame = tk.Frame(win, padx=20, pady=20)
        frame.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(frame, text="Configuração de APIs", font=("Arial", 14, "bold")).pack()
        tk.Label(frame, text="Preencha apenas as APIs que deseja usar", font=("Arial", 9)).pack(pady=5)
        
        # Campos
        campos = [
            ("Google (Gemini):", "api_google"),
            ("Anthropic (Claude):", "api_anthropic"),
            ("OpenAI (GPT):", "api_openai"),
            ("xAI (Grok):", "api_xai"),
        ]
        
        entries = {}
        for label, attr in campos:
            f = tk.Frame(frame)
            f.pack(fill=tk.X, pady=5)
            tk.Label(f, text=label, width=18, anchor="e").pack(side=tk.LEFT)
            e = tk.Entry(f, width=45, show="*")
            e.insert(0, getattr(self.config, attr, ''))
            e.pack(side=tk.LEFT, padx=5)
            entries[attr] = e
        
        # Ollama
        tk.Label(frame, text="─" * 40).pack(pady=10)
        tk.Label(frame, text="Configuração Local (Ollama)", font=("Arial", 11, "bold")).pack()
        
        f = tk.Frame(frame)
        f.pack(fill=tk.X, pady=5)
        tk.Label(f, text="Host:", width=18, anchor="e").pack(side=tk.LEFT)
        e_host = tk.Entry(f, width=45)
        e_host.insert(0, self.config.ollama_host)
        e_host.pack(side=tk.LEFT, padx=5)
        
        f = tk.Frame(frame)
        f.pack(fill=tk.X, pady=5)
        tk.Label(f, text="Modelo:", width=18, anchor="e").pack(side=tk.LEFT)
        e_modelo = tk.Entry(f, width=45)
        e_modelo.insert(0, self.config.modelo_local)
        e_modelo.pack(side=tk.LEFT, padx=5)
        
        def salvar():
            for attr, entry in entries.items():
                setattr(self.config, attr, entry.get().strip())
            self.config.ollama_host = e_host.get().strip()
            self.config.modelo_local = e_modelo.get().strip()
            salvar_config(self.config, self.pasta_script)
            messagebox.showinfo("Salvo", "Configurações salvas!")
            win.destroy()
        
        tk.Button(frame, text="💾 Salvar", command=salvar, font=("Arial", 11), padx=20).pack(pady=15)
    
    def iniciar_processamento(self):
        if self.processando:
            return
        
        modo = self.modo_var.get()
        
        # Verifica se API está configurada
        if modo == "google" and not self.config.api_google:
            messagebox.showerror("Erro", "API do Google não configurada!\nClique em 'Configurar APIs'.")
            return
        elif modo == "anthropic" and not self.config.api_anthropic:
            messagebox.showerror("Erro", "API da Anthropic não configurada!")
            return
        elif modo == "openai" and not self.config.api_openai:
            messagebox.showerror("Erro", "API da OpenAI não configurada!")
            return
        elif modo == "xai" and not self.config.api_xai:
            messagebox.showerror("Erro", "API da xAI não configurada!")
            return
        
        self.processando = True
        self.btn_processar.config(state=tk.DISABLED)
        self.btn_pasta.config(state=tk.DISABLED)
        self.lbl_status.config(text="Processando...")
        
        # Salva modo como padrão
        self.config.modo_padrao = modo
        salvar_config(self.config, self.pasta_script)
        
        thread = threading.Thread(target=self.processar, args=(modo,))
        thread.start()
    
    def processar(self, modo):
        try:
            self.log(f"\n{'='*50}")
            self.log(f"Iniciando processamento em modo {modo.upper()}")
            self.log(f"{'='*50}\n")
            
            # Verifica Ollama se modo local
            if modo == "local":
                self.log("🔌 Verificando Ollama...")
                try:
                    r = requests.get(f"{self.config.ollama_host}/api/tags", timeout=5)
                    if r.status_code != 200:
                        raise Exception("Ollama não respondeu")
                    self.log("✅ Ollama conectado")
                except:
                    self.log("❌ Ollama não está rodando!")
                    self.log("Execute no WSL2: ollama serve")
                    messagebox.showerror("Erro", "Ollama não está rodando!")
                    return
            
            # Processa
            resultado = processar_processo(
                self.pasta_selecionada, modo, self.config, self.log
            )
            
            if not resultado['dados'].numero and not resultado['extracao']:
                messagebox.showwarning("Aviso", "Não foi possível extrair dados do processo!")
                return
            
            # Gera relatórios
            self.log("\n📝 Gerando relatórios...")
            
            # Markdown (com BOM para melhor compatibilidade Windows)
            md_content = gerar_markdown(resultado, self.pasta_selecionada)
            md_path = self.pasta_selecionada / "sintese_processual.md"
            with open(md_path, 'w', encoding='utf-8-sig') as f:
                f.write(md_content)
            self.log(f"  ✅ {md_path.name}")
            
            # Word
            docx_doc = gerar_docx(resultado, self.pasta_selecionada)
            docx_path = self.pasta_selecionada / "sintese_processual.docx"
            docx_doc.save(docx_path)
            self.log(f"  ✅ {docx_path.name}")
            
            # Resultado
            tempo = resultado.get('tempo', 0)
            self.log(f"\n✅ Concluído em {tempo:.1f} segundos!")
            
            messagebox.showinfo(
                "Concluído!",
                f"Síntese gerada em {tempo:.1f} segundos!\n\n"
                f"Arquivos salvos em:\n{self.pasta_selecionada}"
            )
            
            os.startfile(self.pasta_selecionada)
            
        except Exception as e:
            self.log(f"\n❌ Erro: {e}")
            import traceback
            self.log(traceback.format_exc())
            messagebox.showerror("Erro", str(e))
        
        finally:
            self.processando = False
            self.root.after(0, lambda: self.btn_processar.config(state=tk.NORMAL))
            self.root.after(0, lambda: self.btn_pasta.config(state=tk.NORMAL))
            self.root.after(0, lambda: self.lbl_status.config(text="Pronto"))
    
    def executar(self):
        self.root.mainloop()


# ============================================================================
# MAIN
# ============================================================================

if __name__ == "__main__":
    if len(sys.argv) > 1:
        # Modo CLI
        pasta = Path(sys.argv[1])
        modo = sys.argv[2] if len(sys.argv) > 2 else "local"
        
        if not pasta.exists():
            print(f"Erro: Pasta não encontrada: {pasta}")
            sys.exit(1)
        
        config = carregar_config(Path(__file__).parent)
        resultado = processar_processo(pasta, modo, config)
        
        if resultado['dados'].numero or resultado['extracao']:
            md = gerar_markdown(resultado, pasta)
            (pasta / "sintese_processual.md").write_text(md, encoding='utf-8')
            print(f"\n✅ Síntese salva em: {pasta / 'sintese_processual.md'}")
    else:
        # Modo GUI
        gui = BotSinteseGUI()
        gui.executar()
